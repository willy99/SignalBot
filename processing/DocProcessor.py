from docx import Document
import textract
import fitz
from pathlib import Path

import config
from utils.utils import clean_text, format_to_excel_date, get_file_name
import dics.deserter_xls_dic as col
from dics.deserter_xls_dic import *
from datetime import datetime
import re

class DocProcessor:

    NA = ''
    PIECE_HEADER = 'header'
    PIECE_1 = 'piece 1'
    PIECE_2 = 'piece 2'
    PIECE_3 = 'piece 3'
    PIECE_4 = 'piece 4'

    def __init__(self, workflow, file_path):
        self.file_path = file_path
        self.workflow = workflow
        self.response = {
            'insertionDate' :None,
        }

    def process(self):
        extension = Path(self.file_path).suffix
        print("Пошук тексту..." + extension)

        if extension.lower() == '.doc':
            text_pieces = self.get_doc_pieces()
            result = self.process_fields(text_pieces)
            self.workflow.stats.attachmentWordProcessed += 1
            self.workflow.stats.doc_names.append(self.file_path)
            return result
        elif extension.lower() == '.docx':
            text_pieces = self.get_docx_pieces()
            result = self.process_fields(text_pieces)
            self.workflow.stats.attachmentWordProcessed += 1
            self.workflow.stats.doc_names.append(self.file_path)
            return result
        elif extension.lower() == '.pdf':
            self.workflow.stats.attachmentPDFProcessed += 1
            return None
        print("...Пошук закінчено")


    def get_doc_pieces(self):
        doc_pieces = {}
        doc_pieces[self.PIECE_HEADER] = clean_text(self.find_paragraph_doc(TEXT_ANCHOR_HEADER))
        doc_pieces[self.PIECE_1] = clean_text(self.extract_text_between_doc(TEXT_ANCHOR_PIECE_1_START_V1, TEXT_ANCHOR_PIECE_1_END_V1))
        doc_pieces[self.PIECE_3] = clean_text(self.extract_text_between_doc(TEXT_ANCHOR_PIECE_3_START_V1, TEXT_ANCHOR_PIECE_3_END_V1)) or clean_text(
            self.extract_text_between_doc(TEXT_ANCHOR_PIECE_3_START_V2, TEXT_ANCHOR_PIECE_3_END_V1)) or clean_text(
            self.extract_text_between_doc(TEXT_ANCHOR_PIECE_3_START_V3, TEXT_ANCHOR_PIECE_3_END_V1))
        doc_pieces[self.PIECE_4] = clean_text(self.extract_text_between_doc(TEXT_ANCHOR_PIECE_4_START_V1, TEXT_ANCHOR_PIECE_4_END_V1))
        return doc_pieces

    def get_docx_pieces(self):
        doc_pieces = {}
        doc_pieces[self.PIECE_HEADER] = clean_text(self.find_paragraph_docx(TEXT_ANCHOR_HEADER))
        doc_pieces[self.PIECE_1] = clean_text(self.extract_text_between_docx(TEXT_ANCHOR_PIECE_1_START_V1, TEXT_ANCHOR_PIECE_1_END_V1))
        doc_pieces[self.PIECE_3] = clean_text(self.extract_text_between_docx(TEXT_ANCHOR_PIECE_3_START_V1, TEXT_ANCHOR_PIECE_3_END_V1)) or clean_text(
            self.extract_text_between_docx(TEXT_ANCHOR_PIECE_3_START_V2, TEXT_ANCHOR_PIECE_3_END_V1)) or clean_text(
            self.extract_text_between_docx(TEXT_ANCHOR_PIECE_3_START_V3, TEXT_ANCHOR_PIECE_3_END_V1))
        doc_pieces[self.PIECE_4] = clean_text(self.extract_text_between_docx(TEXT_ANCHOR_PIECE_4_START_V1, TEXT_ANCHOR_PIECE_4_END_V1))
        return doc_pieces

    def process_fields(self, text_pieces):

        result = []

        fields = {
            col.COLUMN_INSERT_DATE: format_to_excel_date(datetime.now()),
            col.COLUMN_MIL_UNIT: "А0224",  # Значення за замовчуванням
        }

        text = text_pieces[self.PIECE_HEADER]
        if text is not None:
            fields[col.COLUMN_MIL_UNIT] = self._extract_mil_unit(text)

        text = text_pieces[self.PIECE_1]
        if text is not None:
            # Приклад наповнення результатів після аналізу тексту
            fields[col.COLUMN_DESERTION_DATE] = self._extract_desertion_date(text)
            fields[col.COLUMN_DESERTION_REGION] = self._extract_desertion_region(text)
            fields[col.COLUMN_DESERT_CONDITIONS] = self._extract_desert_conditions(text)
            fields[col.COLUMN_DESERTION_PLACE] = self._extract_desertion_place(text, get_file_name(self.file_path))
            fields[col.COLUMN_RETURN_DATE] = self._extract_return_date(text)

        text = text_pieces[self.PIECE_3]
        if text is not None:
            # Приклад наповнення результатів після аналізу тексту
            fields[col.COLUMN_NAME] = self._extract_name(text)
            fields[col.COLUMN_ID_NUMBER] = self._extract_id_number(text)
            fields[col.COLUMN_TZK] = self._extract_rtzk(text)
            fields[col.COLUMN_PHONE] = self._extract_phone(text)
            fields[col.COLUMN_BIRTHDAY] = self._extract_birthday(text)
            fields[col.COLUMN_TITLE] = self._extract_title(text)
            fields[col.COLUMN_SERVICE_TYPE] = self._extract_service_type(text)
            fields[col.COLUMN_ADDRESS] = self._extract_address(text)
            fields[col.COLUMN_BIO] = self._extract_bio(text, fields[col.COLUMN_NAME])
            fields[col.COLUMN_ENLISTMENT_DATE] = self._extract_conscription_date(text)
            fields[col.COLUMN_SUBUNIT] = self._extract_military_subunit(text, get_file_name(self.file_path))

        fields[col.COLUMN_SERVICE_DAYS] = self._calculate_service_days(fields[col.COLUMN_ENLISTMENT_DATE], fields[col.COLUMN_DESERTION_DATE])

        text = text_pieces[self.PIECE_4]
        fields[col.COLUMN_EXECUTOR] = self._extract_name(text)
        result.append(fields)
        return result

    def _extract_mil_unit(self, text):
        pattern = r'\b[А-ЯA-Z]\d{4}\b'
        match = re.search(pattern, text)
        if match:
            return match.group(0).upper()
        return self.NA

    def _extract_bio(self, text, full_name):
        """
        Повертає частину тексту, починаючи з ПІБ.
        Оскільки текст вичищений, використовуємо прямий пошук.
        """
        if not full_name or full_name == self.NA:
            return text

        # Знаходимо позицію, де починається ПІБ
        start_index = text.find(full_name)

        # Якщо знайшли — ріжемо, якщо ні — повертаємо як є
        if start_index != -1:
            return text[start_index:].strip()

        return text

    def _extract_name(self, text):
        # беремо саме ПЕРШИЙ знайдений ПІБ у тексті
        pattern = r'\b([А-ЩЬЮЯҐЄІЇ]{3,})\s+([А-ЯҐЄІЇ][а-яґєії\']{2,})\s+([А-ЯҐЄІЇ][а-яґєії\']{2,})\b'

        match = re.search(pattern, text)
        if match:
            return f"{match.group(1)} {match.group(2)} {match.group(3)}"
        return self.NA

    def _extract_title(self, text):
        """
        Шукає військове звання зі списку ключових слів.
        """

        # Створюємо regex з переліку звань
        pattern = r'\b(' + '|'.join(TITLES) + r')\b'

        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0).lower()
        return self.NA

    def _extract_service_type(self, text):
        mapping = {
            r'приз[ио]вом|мобілізаці': "призовом",
            r'контрактом': "контрактом"
        }
        for pattern, result in mapping.items():
            if re.search(pattern, text, re.IGNORECASE):
                return result
        return self.NA

    def _extract_id_number(self, text):
        """
        Шукає 10 цифр підряд поруч із ключовими словами РНОКПП або ІПН.
        Враховує можливі переплутані літери (напр. РНОК ПП, І.П.Н).
        """
        # Шукаємо ключові слова, після яких йде 10 цифр
        # [РІ][НП][ОН][К][П]? - спроба охопити помилки в РНОКПП/ІПН
        pattern = r'(?:РНОКПП|ІПН|І\.П\.Н\.|РНОК\s*ПП|РНОК\s*ПП)[\s.:]*(\d{10})\b'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)

        # Якщо поруч немає ключових слів, шукаємо просто 10 цифр,
        # але перевіряємо, щоб це не був номер телефону
        standalone_digits = re.findall(r'\b\d{10}\b', text)
        return standalone_digits[0] if standalone_digits else self.NA

    def _extract_phone(self, text):
        """
        Шукає номер телефону та приводить його до формату 0991141111.
        """
        # Шукаємо цифри, що йдуть після слова "номер" або "телефону"
        # Паттерн охоплює +38, 8, дужки та дефіси
        pattern = r'(?:номер|тел)[\s\w.:]*(\+?3?8?[\s(-]*0\d{2}[\s)-]*\d{3}[\s-]*\d{2}[\s-]*\d{2})\b'
        match = re.search(pattern, text, re.IGNORECASE)

        if match:
            raw_phone = match.group(1)
            # Очищаємо від усього, крім цифр
            digits = re.sub(r'\D', '', raw_phone)
            # Приводимо до формату 0XXXXXXXXX (останні 10 цифр)
            if len(digits) >= 10:
                return digits[-10:]
        return self.NA

    def _extract_conscription_date(self, text):
        """
        Витягує дату призову. Підтримує формати: 01.01.2025 та 01.01.25.
        """
        # Патерн для дати: ДД.ММ.РРРР або ДД.ММ.РР
        # \d{2} - день, \d{2} - місяць, (\d{4}|\d{2}) - рік (4 або 2 цифри)
        date_pattern = r'(\d{2}\.\d{2}\.(?:\d{4}|\d{2}))'

        # 1. Знаходимо РНОКПП/ІПН, щоб знати, де зупинити пошук
        id_match = re.search(r'(?:РНОКПП|ІПН)', text, re.IGNORECASE)
        if not id_match:
            return self.NA

        pos_id = id_match.start()
        # Беремо зону пошуку ПЕРЕД РНОКПП (близько 150 символів)
        lookback_area = text[max(0, pos_id - 150):pos_id]

        # 2. Шукаємо всі дати в цій зоні
        dates = re.findall(date_pattern, lookback_area)

        if dates:
            # Беремо останню дату (найближчу до РНОКПП)
            found_date = dates[-1]
            return format_to_excel_date(found_date)

        return self.NA

    def _extract_birthday(self, text):
        """
        Витягує дату перед словами "року народження".
        """
        # Шукаємо формат ДД.ММ.РРРР перед "року народження"
        pattern = r'(\d{2}\.\d{2}\.\d{4})[\s]*(року народження|р.н.)'
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return format_to_excel_date(match.group(1))
        return self.NA

    def _extract_address(self, text):
        """
        Беремо текст після 'Адреса проживання' і шукаємо в ньому
        найдовшу послідовність, що закінчується номером будинку/квартири.
        """
        # 1. Знаходимо ОСТАННЄ входження ключової фрази
        marker = "Адреса проживання"
        starts = [m.start() for m in re.finditer(re.escape(marker), text, re.IGNORECASE)]
        if not starts:
            return self.NA

        # Працюємо з текстом після маркера
        search_area = text[starts[-1]:]
        # Видаляємо сам заголовок
        search_area = re.sub(r'^Адреса проживання(?: військовослужбовця)?\s*:?\s*', '', search_area,
                             flags=re.IGNORECASE)

        # 2. ПАТТЕРН: Беремо все, поки не зустрінемо номер будинку
        # (цифри, можливо з дробом або літерою), і опціонально квартиру.
        # [^.!?\n]{5,} - мінімум 5 будь-яких символів (вулиця, місто)
        # \d{1,4}(?:[/-]\d+)?\s*[А-Яа-я]? - номер будинку
        pattern = r'([^.!?\n]{5,}.*?\d{1,4}(?:[/-]\d+)?\s*[А-Яа-я]?(?:\s*,?\s*кв\.?\s*\d+)?)'

        match = re.search(pattern, search_area, re.DOTALL)

        if match:
            address = match.group(1).strip()

            # Видаляємо "хвости", які могли прилипнути (наприклад, початок наступного пункту)
            # Якщо в адресі раптом з'явився текст "4. Хто проводив" - обрізаємо
            address = re.split(r'\s\d+\s*[А-Я]', address)[0]

            # Фінальна чистка пробілів
            return " ".join(address.split()).strip(':;,. ')

        return self.NA

    def _extract_rtzk(self, text):
        pattern = r'(?i)([А-ЯҐЄІЇ][^.,!?]*?(?:РТЦК|ТЦК|МТЦК)(?:\s*(?:та|&)?\s*СП)?(?:\s+м\.\s+[А-Яа-я\']+|\s+[А-Яа-я\']+\s+обл\.?)?)'

        match = re.search(pattern, text)
        if match:
            res = match.group(1).strip()

            res = re.sub(
                r'(?i)^(Призваний|Призвана|Яким)\s+(на військову службу\s+)?(за призовом\s+)?(під час мобілізації\s+)?',
                '', res)

            res = re.sub(r'\d{2}\.\d{2}\.\d{4}.*$', '', res)

            final_res = " ".join(res.split()).strip(':;,. ')
            return final_res if final_res else self.NA

        return self.NA

    def _extract_desertion_date(self, text):
        """
        Шукає дату, пов'язану з моментом зникнення (під час перевірки або залишення).
        """
        # Шукаємо дату ДД.ММ.РРРР, яка стоїть перед описом відсутності
        pattern = r'(\d{2}\.\d{2}\.\d{4})(?=\s+року\s+(?:під час перевірки|був відсутній|самовільно залишив))'
        match = re.search(pattern, text, re.IGNORECASE)

        if match:
            # Використовуємо вашу функцію форматування для отримання m/d/yy
            return format_to_excel_date(match.group(1))

        # Якщо за специфічним якорем не знайдено, беремо першу дату в блоці обставин
        fallback = re.search(r'(\d{2}\.\d{2}\.\d{4})', text)
        if fallback:
            return format_to_excel_date(fallback.group(1))

        return self.NA

    def _extract_desert_conditions(self, text):
        """
        Шукає абзац з обставинами (перевірка наявності/шикування + відсутність).
        """
        # 1. Універсальне розбиття на абзаци
        paragraphs = [p.strip() for p in re.split(r'[\r\n]{2,}', text) if p.strip()]

        # Ключові слова для перевірки/контролю
        check_markers = ["під час перевірки", "під час шикування", "перевірці наявності", "не повернувся"]
        # Ключові слова для факту відсутності
        absence_markers = ["відсутн", "виявлено відсутність", "не було в наявності", "не повернувся"]

        for para in paragraphs:
            clean_para = " ".join(para.split()).lower()

            # 2. Перевіряємо, чи є хоча б один маркер перевірки ТА хоча б один маркер відсутності
            has_check = any(marker in clean_para for marker in check_markers)
            has_absence = any(marker in clean_para for marker in absence_markers)

            if has_check and has_absence:
                # Повертаємо абзац одним рядком без внутрішніх розривів
                return " ".join(para.split())

        return self.NA

    def _extract_return_date(self, text):
        """
        Повертає дату тільки якщо в тексті є факт присутності.
        Якщо 'був присутній' не знайдено — повертає N/A.
        """
        # 1. Перевіряємо наявність ключової фрази
        if "був присутній" not in text.lower():
            return self.NA

        # 2. Якщо фраза є, шукаємо дату ДД.ММ.РРРР, яка стоїть ПЕРЕД "року був присутній"
        # Або просто дату в цьому ж реченні
        pattern = r'(\d{2}\.\d{2}\.\d{4})(?=\s+року\s+був\s+присутній)'
        match = re.search(pattern, text, re.IGNORECASE)

        if match:
            return format_to_excel_date(match.group(1))

        # 3. Резервний пошук дати ТІЛЬКИ якщо ми вже знаємо, що людина була присутня
        # (на випадок іншого порядку слів)
        fallback_with_presence = re.search(r'(\d{2}\.\d{2}\.\d{4})', text)
        if fallback_with_presence:
            return format_to_excel_date(fallback_with_presence.group(1))

        return self.NA

    def _extract_desertion_region(self, text):
        """
        Витягує повну географічну назву: населений пункт, район та область.
        """
        # 1. Визначаємо ключові маркери початку (н.п., с., м., місто, село тощо)
        # 2. Захоплюємо все до слова "області" або "обл." включно.
        # Патерн пояснення:
        # (?i) - ігнорувати регістр
        # (?:н\.п\.|с\.|м\.|село|місто|селище|смт)\s+ - початок з маркера
        # ([А-ЯҐЄІЇ].*?(?:області|обл\.)) - назва з великої літери до слова область

        pattern = r'(?i)(?:н\.п\.|с\.|м\.|село|місто|селище|смт)\s+([А-ЯҐЄІЇ][^.;]*?(?:області|обл\.))'

        match = re.search(pattern, text, re.DOTALL)
        if match:
            # Очищаємо від зайвих пробілів та переносів рядків
            full_address = " ".join(match.group(1).split())
            # Видаляємо зайві крапки в кінці, якщо вони є
            return full_address.strip().rstrip('.')

        # Додатковий пошук: якщо не знайшли маркер н.п., шукаємо просто конструкцію "Район... Область"
        backup_pattern = r'(?i)([А-Я][а-яіЇє]*?\s+район[у|а]\s+[А-Я][а-яіЇє]*?\s+області)'
        backup_match = re.search(backup_pattern, text)
        if backup_match:
            return " ".join(backup_match.group(1).split())

        return self.NA

    def _calculate_service_days(self, conscription_date_str, desertion_date_str):
        """
            Рахує дні та перевіряє їх на логіку.
            """
        if conscription_date_str == self.NA or desertion_date_str == self.NA:
            return 0

        try:
            def parse_date(d_str):
                return datetime.strptime(d_str, config.EXCEL_DATE_FORMAT)

            dt_start = parse_date(conscription_date_str)
            dt_end = parse_date(desertion_date_str)

            # Рахуємо різницю
            delta = dt_end - dt_start
            days = delta.days

            if days < 0 or days > 4000:
                print(f"[УВАГА] Нелогічна кількість днів служби: {days}. Перевірте дати.")
                return 0

            return days

        except Exception as e:
            print(f"Помилка розрахунку днів: {e}")
            return 0

    def _extract_military_subunit(self, text, file_name=None):
        # 1. Автоматично формуємо список унікальних значень із мапінгу (Values)
        short_values = set()
        for val in SUBUNIT_MAPPING.values():
            # Видаляємо \1, щоб отримати чисту абревіатуру для пошуку
            clean_val = val.replace(r'\1', '').strip()
            if clean_val:
                short_values.add(clean_val)

        # ЕТАП 1: Перевірка назви файлу (тепер з IGNORECASE)
        if file_name:
            # Сортуємо від найдовших до найкоротших
            sorted_shorts = sorted(short_values, key=len, reverse=True)

            for short_val in sorted_shorts:
                pattern = rf'(?:^|[\s_])(\d*[\s_]*)?{re.escape(short_val)}(?=[\s_]|$)'

                # Додаємо re.IGNORECASE, щоб 'аемб' або 'Аемб' теж працювали
                match = re.search(pattern, file_name, re.IGNORECASE)

                if match:
                    res = match.group(0).strip()
                    # Гарне форматування: додаємо пробіл між цифрою та текстом, якщо його не було
                    return re.sub(rf'(\d+)\s*({re.escape(short_val)})', r'\1 \2', res, flags=re.IGNORECASE).replace('_','')

        # ЕТАП 2: Пошук у тексті через мапінг (якщо в файлі не знайдено)
        found_subunits = []
        for pattern, abbreviation in SUBUNIT_MAPPING.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                if r'\1' in abbreviation:
                    digit = match.group(1) if match.group(1) else ""
                    res = abbreviation.replace(r'\1', digit).strip()
                else:
                    res = abbreviation
                found_subunits.append(res)

        return found_subunits[-1] if found_subunits else self.NA

    def _extract_desertion_place(self, text, file_name=None):
        short_values = list(set(DESERTION_PLACE_MAPPING.values()))

        if file_name:
            for val in short_values:
                if re.search(rf'\b{re.escape(val)}\b', file_name, re.IGNORECASE):
                    return val

        for pattern, short_name in DESERTION_PLACE_MAPPING.items():
            if re.search(pattern, text, re.IGNORECASE):
                return short_name

        return self.NA


    ########################### - DOC - ###############################
    def find_paragraph_doc(self, search_text, get_next=False):
        # textract витягує текст з .doc через antiword
        byte_content = textract.process(self.file_path)
        text = byte_content.decode('utf-8')

        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

        for i, para in enumerate(paragraphs):
            para = clean_text(para).lower()
            if search_text.lower() in para:
                if get_next and i + 1 < len(paragraphs):
                    return clean_text(paragraphs[i + 1])
                else:
                    return clean_text(paragraphs[i])
        return None

    def extract_text_between_doc(self, start_search, end_search):
        # 1. Отримуємо весь текст одним шматком
        byte_content = textract.process(self.file_path)
        # Замінюємо декілька пробілів/переносів на один пробіл для стабільності пошуку
        full_text = " ".join(byte_content.decode('utf-8').split())

        # 2. Готуємо пошукові фрази (теж очищаємо від зайвих пробілів)
        start_phrase = " ".join(start_search.split())
        end_phrase = " ".join(end_search.split())

        # 3. Шукаємо позицію кінця стартової фрази (rfind знайде ОСТАННЄ входження, тобто після інструкції)
        start_pos = full_text.lower().rfind(start_phrase.lower())

        if start_pos == -1:
            return None

        # Початок контенту — одразу після стартової фрази
        content_start = start_pos + len(start_phrase)

        # 4. Шукаємо позицію кінцевої фрази (починаючи від content_start)
        end_pos = full_text.lower().find(end_phrase.lower(), content_start)

        # 5. Вирізаємо результат
        if end_pos != -1:
            result = full_text[content_start:end_pos].strip()
        else:
            result = full_text[content_start:].strip()

        return result if result else None


    ########################### - DOCX - ###############################
    def find_paragraph_docx(self, search_text, get_next=False):
        """
        Аналог для .docx файлів. Шукає параграф за текстом
        і повертає його або наступний за ним.
        """
        doc = Document(self.file_path)
        # Створюємо список параграфів, які не є порожніми
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        for i, para_text in enumerate(paragraphs):
            # Використовуємо clean_text для нормалізації пошуку (як у вашому прикладі)
            para_clean = clean_text(para_text).lower()

            if search_text.lower() in para_clean:
                if get_next and i + 1 < len(paragraphs):
                    return clean_text(paragraphs[i + 1])
                else:
                    return clean_text(paragraphs[i])

        return None

    def extract_text_between_docx(self, start_search, end_search):
        """
        Аналог для .docx. Збирає весь текст документа в один рядок
        і вирізає контент між двома фразами.
        """
        # 1. Завантажуємо документ
        doc = Document(self.file_path)

        # 2. Збираємо весь текст, нормалізуючи пробіли (як у методі для .doc)
        # Це дозволяє ігнорувати розриви сторінок та параграфів під час пошуку
        full_text = " ".join([p.text for p in doc.paragraphs]).strip()
        full_text = " ".join(full_text.split())

        # 3. Готуємо пошукові фрази
        start_phrase = " ".join(start_search.split())
        end_phrase = " ".join(end_search.split())

        # 4. Шукаємо останнє входження стартової фрази (rfind)
        start_pos = full_text.lower().rfind(start_phrase.lower())

        if start_pos == -1:
            return None

        # Початок контенту — одразу після знайденої фрази
        content_start = start_pos + len(start_phrase)

        # 5. Шукаємо позицію кінцевої фрази
        end_pos = full_text.lower().find(end_phrase.lower(), content_start)

        # 6. Вирізаємо та чистимо результат
        if end_pos != -1:
            result = full_text[content_start:end_pos].strip()
        else:
            result = full_text[content_start:].strip()

        return result if result else None
    ########################### - PDF - ###############################
    def find_next_paragraph_pdf(self, search_text):
        print('>>> pdf')
        doc = fitz.open(self.file_path)

        for page in doc:
            # Отримуємо блоки тексту. Кожен блок зазвичай є абзацом.
            blocks = page.get_text("blocks")
            for i, b in enumerate(blocks):
                block_text = b[4]  # 4-й елемент кортежу — це сам текст
                print('>>> block : ' + clean_text(block_text))
                if search_text.lower() in clean_text(block_text.lower()):
                    if i + 1 < len(blocks):
                        return clean_text(blocks[i + 1][4])
                    return "Знайдено в останньому блоці сторінки."

        return "Текст не знайдено."

