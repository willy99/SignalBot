from docx import Document
import textract
import fitz

def find_next_paragraph_doc(file_path, search_text):
    print('>>> doc')
    # textract витягує текст з .doc через antiword
    byte_content = textract.process(file_path)
    text = byte_content.decode('utf-8')

    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]

    for i, para in enumerate(paragraphs):
        para = clean_text(para).lower()
        if search_text.lower() in para:
            if i + 1 < len(paragraphs):
                return clean_text(paragraphs[i + 1])
    return "Не знайдено"

def find_next_paragraph_docx(file_path, search_text):
    print('>>> docx')
    doc = Document(file_path)
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in clean_text(para.text.lower()):
            # Перевіряємо, чи є наступний абзац
            if i + 1 < len(doc.paragraphs):
                return clean_text(doc.paragraphs[i+1].text)
            return "Це був останній абзац."
    return "Текст не знайдено."


def find_next_paragraph_pdf(file_path, search_text):
    print('>>> pdf')
    doc = fitz.open(file_path)

    for page in doc:
        # Отримуємо блоки тексту. Кожен блок зазвичай є абзацом.
        blocks = page.get_text("blocks")
        for i, b in enumerate(blocks):
            block_text = b[4]  # 4-й елемент кортежу — це сам текст
            print('>>> block : ' + clean_text(block_text))
            if search_text.lower() in clean_text(block_text.lower()):
                if i + 1 < len(blocks):
                    return clean_text(blocks[i + 1][4])
                return "Знайдено в останньому блоці сторінки."

    return "Текст не знайдено."

def clean_text(text):
    # .split() без аргументів розбиває рядок по будь-якій кількості
    # пробілів, табуляцій та переносів, а " ".join зшиває їх назад одним пробілом.
    return " ".join(text.split())
